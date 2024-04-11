# Written by Mark Lindberg. Copyright 2017-2023.
# Mucho thanks to Ravi and Ross for testing.
# License: Modify and use as you wish.

import zipfile
import re
from sys import argv
from lxml import etree
from argparse import ArgumentParser
from os.path import isfile

w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
ns = {"w": w}


def insertParNums(old, newMargin, newInline):
    """
    This inserts a textbox to the left of the paragraph, containing the
    paragraph number.
    """
    zin = zipfile.ZipFile(old, "r")
    zoutMargin = zipfile.ZipFile(newInline, "w")
    # .docx files are actually a zipped up set of xml files.
    for item in zin.infolist():
        buf = zin.read(item.filename)
        # This is the one that contains the text of the document.
        if item.filename == "word/document.xml":
            text = buf.decode("utf-8")
            parser = etree.XMLParser(no_network=False)
            root = etree.fromstring(text.encode(), parser)
            paragraphs = root.findall(".//w:p", ns)
            count = 1
            for p in paragraphs:
                # Paragraphs that have text.
                if p.findall(".//w:t", ns):
                    first_run = p.find(".//w:r", ns)
                    first_text = first_run.find(".//w:t", ns)
                    lower_text = first_text.text.lower().strip()
                    # Reset numbering on each chapter.
                    if (
                        lower_text.startswith("prologue")
                        or lower_text.startswith("chapter")
                        or lower_text.startswith("epilogue")
                    ):
                        count = 1
                        continue
                    my_run = etree.Element(f"{{{w}}}r")
                    run_rpr = etree.SubElement(my_run, f"{{{w}}}rPr")
                    run_align = etree.SubElement(run_rpr, f"{{{w}}}vertAlign")
                    run_align.set(f"{{{w}}}val", "superscript")
                    run_color = etree.SubElement(run_rpr, f"{{{w}}}color")
                    run_color.set(f"{{{w}}}val", "A7A7A7")
                    run_text = etree.SubElement(my_run, f"{{{w}}}t")
                    run_text.text = str(count)
                    first_run.addprevious(my_run)
                    count += 1
            # Write out all files, including modified text.
            zoutMargin.writestr(item, etree.tostring(root))
        else:
            zoutMargin.writestr(item, buf)
    zoutMargin.close()
    zin.close()


def docxParNums(oldFile, newFile):
    """
    This will insert the paragraph numbers using the docx library, but also has
    to copy over the text of the paragraph, and may lose formatting.
    """
    from docx import Document
    from docx.shared import RGBColor, Inches

    doc = Document(oldFile)
    count = 1
    gray = RGBColor(0xA7, 0xA7, 0xA7)
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "":
            continue
        lower_text = paragraph.text.lower().strip()
        if (
            lower_text.startswith("prologue")
            or lower_text.startswith("chapter")
            or lower_text.startswith("epilogue")
        ):
            count = 1
            continue
        runs = paragraph.runs
        paragraph.clear()
        paragraph.paragraph_format.first_line_indent = 0
        paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(0.25))
        num_run = paragraph.add_run(str(count))
        num_run.font.superscript = True
        num_run.font.color.rgb = gray
        paragraph.add_run("\t")  # Indent.
        for run in runs:
            new_run = paragraph.add_run(run.text)
            new_run.style = run.style
            new_run.underline = run.underline
            new_run.italic = run.italic
            new_run.bold = run.bold
            # new_run.font = run.font
        count += 1
    doc.save(newFile)


def viewText(fname):
    """
    Print the entire underlying xml file.
    """
    zin = zipfile.ZipFile(fname, "r")
    for item in zin.infolist():
        if item.filename == "word/document.xml":
            text = zin.read(item.filename).decode("utf-8")
            pr = True
            count = 0
            for line in text.split("</w:p>"):
                count += 1
                if "EPILOGUE" in line:
                    pr = True
                if pr:
                    print()
                    print(line)
                if count > 10:
                    break
    zin.close()


if __name__ == "__main__":
    parser = ArgumentParser()
    parser.add_argument(
        "--view",
        action="store_true",
        help="Print the raw xml without modifying the file.",
    )
    parser.add_argument(
        "--docx", action="store_true", help="Run the python-docx version."
    )
    parser.add_argument("filename")
    args = parser.parse_args()
    filename = args.filename
    if not filename.endswith(".docx"):
        filename += ".docx"
    assert isfile(filename), f"{filename} not found."
    marginFile = filename.split(".docx")[0] + " with margin numbers.docx"
    inlineFile = filename.split(".docx")[0] + " with inline numbers.docx"
    if args.view:
        viewText(filename)
    elif args.docx:
        docxParNums(filename, inlineFile)
    else:
        insertParNums(filename, marginFile, inlineFile)
